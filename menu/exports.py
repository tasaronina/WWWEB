# Экспорт заказов. Только для админов. Исправлено под поля qty и menu.price.

from datetime import datetime
from decimal import Decimal
import io

from django.db.models import Q
from django.http import HttpResponse
from rest_framework.permissions import IsAuthenticated
from rest_framework.views import APIView

from openpyxl import Workbook
from openpyxl.utils import get_column_letter
from docx import Document

from .models import Order
from .permissions import IsAdmin


STATUS_LABELS = {
    "NEW": "Новый",
    "IN_PROGRESS": "Готовится",
    "DONE": "Готов",
    "PAID": "Оплачен",
    "CANCELLED": "Отменён",
    "CANCELED": "Отменён",
    "CANCEL": "Отменён",
}


def _iter_items(order):
    for rel in ("items", "order_items", "orderitem_set"):
        if hasattr(order, rel):
            qs = getattr(order, rel)
            try:
                return qs.all()
            except Exception:
                return qs
    return []


def _customer_name(order):
    c = getattr(order, "customer", None)
    if not c:
        return "—"
    for a in ("name", "full_name", "fio"):
        if hasattr(c, a):
            return getattr(c, a)
    first = getattr(c, "first_name", "")
    last = getattr(c, "last_name", "")
    mid = getattr(c, "middle_name", "")
    t = " ".join(x for x in (last, first, mid) if x)
    return t or f"ID {getattr(c, 'id', '—')}"


def _unit_price(item):
    # для твоей модели OrderItem: qty + menu.price
    m = getattr(item, "menu", None)
    if m and hasattr(m, "price") and m.price is not None:
        return Decimal(m.price)
    return Decimal("0")


def _excel_autowidth(ws):
    for col in ws.columns:
        length = 0
        idx = col[0].column if hasattr(col[0], "column") else 1
        for cell in col:
            try:
                length = max(length, len(str(cell.value)))
            except Exception:
                pass
        ws.column_dimensions[get_column_letter(idx)].width = min(length + 2, 60)


def _make_excel(sheets, row_limit=50000):
    wb = Workbook()
    # первый лист
    title, headers, rows = sheets[0]
    ws = wb.active
    ws.title = title[:31] or "Sheet1"
    if headers:
        ws.append(headers)
    for r in rows[:row_limit]:
        ws.append(r)
    _excel_autowidth(ws)

    # остальные
    for title, headers, rows in sheets[1:]:
        ws = wb.create_sheet(title=title[:31] or "Sheet")
        if headers:
            ws.append(headers)
        for r in rows[:row_limit]:
            ws.append(r)
        _excel_autowidth(ws)

    stream = io.BytesIO()
    wb.save(stream)
    stream.seek(0)
    return stream


def _make_word(title, blocks, row_limit=50000):
    doc = Document()
    doc.add_heading(title, 0)
    for blk in blocks:
        t = blk.get("type")
        if t == "text":
            doc.add_paragraph(blk.get("text", ""))
        elif t == "table":
            headers = blk.get("headers") or []
            rows = blk.get("rows") or []
            table = doc.add_table(rows=1, cols=len(headers))
            hdr = table.rows[0].cells
            for i, h in enumerate(headers):
                hdr[i].text = str(h)
            for r in rows[:row_limit]:
                cells = table.add_row().cells
                for j, v in enumerate(r):
                    cells[j].text = "" if v is None else str(v)
    stream = io.BytesIO()
    doc.save(stream)
    stream.seek(0)
    return stream


class OrdersExportView(APIView):
    permission_classes = [IsAuthenticated, IsAdmin]

    def get_queryset(self, request):
        qs = Order.objects.select_related("customer", "user").prefetch_related("items__menu").order_by("-id")
        # Допфильтры для админского экспорта
        p = request.query_params
        status = p.get("status")
        search = p.get("search")
        if status:
            qs = qs.filter(status=status)
        if search:
            cond = Q(id__icontains=search)
            try:
                cond |= Q(customer__name__icontains=search)
            except Exception:
                pass
            qs = qs.filter(cond)
        return qs

    def get(self, request):
        qs = self.get_queryset(request)

        file_type = (request.query_params.get("type") or "excel").lower()
        include_items = (request.query_params.get("include_items") or "false").lower() == "true"
        include_summary = (request.query_params.get("include_summary") or "true").lower() != "false"

        orders_headers = ["ID", "Клиент", "Статус", "Создан", "Позиций", "Сумма, ₽"]
        orders_rows = []
        for o in qs:
            created = getattr(o, "created_at", None)
            created_str = created.strftime("%Y-%m-%d %H:%M") if created else ""
            items = list(_iter_items(o))
            items_count = len(items)
            total = Decimal("0")
            for it in items:
                qty = int(getattr(it, "qty", 0) or 0)
                up = _unit_price(it)
                total += Decimal(qty) * up
            orders_rows.append([
                o.id,
                _customer_name(o),
                STATUS_LABELS.get(getattr(o, "status", ""), getattr(o, "status", "") or "—"),
                created_str,
                items_count,
                f"{total:.2f}",
            ])

        items_headers = ["ID", "OrderID", "Позиция", "Кол-во", "Цена, ₽", "Итого, ₽"]
        items_rows = []
        if include_items:
            for o in qs:
                for it in _iter_items(o):
                    qty = int(getattr(it, "qty", 0) or 0)
                    up = _unit_price(it)
                    line = Decimal(qty) * up
                    title = ""
                    m = getattr(it, "menu", None)
                    if m:
                        title = getattr(m, "name", "") or getattr(m, "title", "") or f"ID {getattr(m, 'id', '')}"
                    items_rows.append([getattr(it, "id", None), o.id, title, qty, f"{up:.2f}", f"{line:.2f}"])

        summary_headers = ["Показатель", "Значение"]
        summary_rows = []
        if include_summary:
            cnt = len(orders_rows)
            total_sum = sum(Decimal(r[5].replace(",", ".").replace(" ", "")) for r in orders_rows) if orders_rows else Decimal("0")
            avg = (total_sum / cnt) if cnt else Decimal("0")
            summary_rows.extend([
                ["Количество заказов", cnt],
                ["Суммарная выручка, ₽", f"{total_sum:.2f}"],
                ["Средний чек, ₽", f"{avg:.2f}"],
            ])

        now = datetime.now().strftime("%Y-%m-%d_%H-%M")
        base = f"orders_{now}"

        if file_type == "word":
            blocks = [
                {"type": "text", "text": "Экспорт заказов (админ)"},
                {"type": "table", "headers": orders_headers, "rows": orders_rows},
            ]
            if include_items:
                blocks.append({"type": "table", "headers": items_headers, "rows": items_rows})
            if include_summary:
                blocks.append({"type": "table", "headers": summary_headers, "rows": summary_rows})
            stream = _make_word("Orders Export", blocks)
            resp = HttpResponse(
                stream,
                content_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            )
            resp["Content-Disposition"] = f'attachment; filename="{base}.docx"'
            return resp

        sheets = [("Orders", orders_headers, orders_rows)]
        if include_items:
            sheets.append(("OrderItems", items_headers, items_rows))
        if include_summary:
            sheets.append(("Summary", summary_headers, summary_rows))
        stream = _make_excel(sheets)
        resp = HttpResponse(
                stream,
                content_type="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
        )
        resp["Content-Disposition"] = f'attachment; filename="{base}.xlsx"'
        return resp
