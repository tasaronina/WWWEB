import pyotp

from django.contrib.auth import authenticate, login, logout
from django.http import HttpResponse
from django.utils.decorators import method_decorator
from django.views.decorators.csrf import ensure_csrf_cookie
from django.db.models import Count, Min, Max, Avg, Sum, F, DecimalField, ExpressionWrapper

from rest_framework import permissions, serializers
from rest_framework.decorators import action
from rest_framework.response import Response
from rest_framework.viewsets import ModelViewSet, GenericViewSet

from .models import Category, Menu, Customer, Order, OrderItem, Profile
from .serializers import (
    CategorySerializer,
    MenuSerializer,
    CustomerSerializer,
    OrderSerializer,
    OrderItemSerializer,
)

try:
    from openpyxl import Workbook
except Exception:
    Workbook = None

try:
    from docx import Document
except Exception:
    Document = None


class LoginSerializer(serializers.Serializer):
    username = serializers.CharField()
    password = serializers.CharField()


class SecondLoginSerializer(serializers.Serializer):
    key = serializers.CharField()


class UserViewSet(GenericViewSet):
    permission_classes = [permissions.AllowAny]
    serializer_class = LoginSerializer

    @method_decorator(ensure_csrf_cookie)
    @action(detail=False, url_path="csrf", methods=["GET"])
    def csrf(self, request, *args, **kwargs):
        return Response({"ok": True})

    def get_serializer_class(self):
        if self.action == "second_login":
            return SecondLoginSerializer
        return LoginSerializer

    @action(detail=False, url_path="info", methods=["GET"])
    def get_info(self, request, *args, **kwargs):
        data = {
            "username": request.user.username if request.user.is_authenticated else "",
            "is_authenticated": request.user.is_authenticated,
            "is_staff": request.user.is_staff if request.user.is_authenticated else False,
        }
        if request.user.is_authenticated:
            data["second"] = request.session.get("second") or False
        return Response(data)

    @action(detail=False, url_path="login", methods=["POST"])
    def login_user(self, request, *args, **kwargs):
        ser = LoginSerializer(data=request.data)
        ser.is_valid(raise_exception=True)

        user = authenticate(
            username=ser.validated_data["username"],
            password=ser.validated_data["password"],
        )
        if user is not None:
            login(request, user)
            request.session["second"] = False
            return Response({"success": True})

        return Response({"success": False})

    @action(detail=False, url_path="logout", methods=["POST"], permission_classes=[permissions.IsAuthenticated])
    def logout_user(self, request, *args, **kwargs):
        logout(request)
        return Response({"status": "success"})

    @action(detail=False, url_path="get-totp", methods=["GET"], permission_classes=[permissions.IsAuthenticated])
    def get_totp(self, request, *args, **kwargs):
        profile, _ = Profile.objects.get_or_create(user=request.user)
        profile.opt_key = pyotp.random_base32()
        profile.save()

        url = pyotp.totp.TOTP(profile.opt_key).provisioning_uri(
            name=request.user.username,
            issuer_name="CafeApp",
        )
        return Response({"url": url})

    @action(detail=False, url_path="second-login", methods=["POST"], permission_classes=[permissions.IsAuthenticated])
    def second_login(self, request, *args, **kwargs):
        ser = SecondLoginSerializer(data=request.data)
        ser.is_valid(raise_exception=True)

        profile, _ = Profile.objects.get_or_create(user=request.user)
        if not profile.opt_key:
            return Response({"success": False})

        t = pyotp.totp.TOTP(profile.opt_key)
        code = ser.validated_data["key"].strip()

        if code == t.now():
            request.session["second"] = True
            return Response({"success": True})

        return Response({"success": False})

class CategoryViewSet(ModelViewSet):
    queryset = Category.objects.all().order_by("-id")
    serializer_class = CategorySerializer
    permission_classes = [permissions.IsAuthenticated]

    def get_queryset(self):
        qs = super().get_queryset()
        name = self.request.query_params.get("name")
        if name:
            qs = qs.filter(name__icontains=name)
        return qs

    @action(detail=False, url_path="stats", methods=["GET"])
    def stats(self, request, *args, **kwargs):
        d = self.get_queryset().aggregate(
            total=Count("id"),
        )
        return Response({"total": d.get("total") or 0})

    @action(detail=False, url_path="export-excel", methods=["GET"])
    def export_excel(self, request, *args, **kwargs):
        if Workbook is None:
            return Response({"detail": "openpyxl не установлен в этом окружении"}, status=500)

        qs = self.get_queryset()
        wb = Workbook()
        ws = wb.active
        ws.title = "Категории"
        ws.append(["id", "name"])

        for c in qs:
            ws.append([c.id, c.name])

        resp = HttpResponse(
            content_type="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet"
        )
        resp["Content-Disposition"] = 'attachment; filename="categories.xlsx"'
        wb.save(resp)
        return resp

    @action(detail=False, url_path="export-word", methods=["GET"])
    def export_word(self, request, *args, **kwargs):
        if Document is None:
            return Response({"detail": "python-docx не установлен в этом окружении"}, status=500)

        qs = self.get_queryset()
        doc = Document()
        doc.add_heading("Категории", level=1)

        table = doc.add_table(rows=1, cols=2)
        hdr = table.rows[0].cells
        hdr[0].text = "id"
        hdr[1].text = "name"

        for c in qs:
            row = table.add_row().cells
            row[0].text = str(c.id)
            row[1].text = c.name

        resp = HttpResponse(
            content_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
        )
        resp["Content-Disposition"] = 'attachment; filename="categories.docx"'
        doc.save(resp)
        return resp

class MenuViewSet(ModelViewSet):
    queryset = Menu.objects.all().order_by("-id")
    serializer_class = MenuSerializer
    permission_classes = [permissions.IsAuthenticated]

    def get_queryset(self):
        qs = Menu.objects.all().order_by("-id")

        title = self.request.query_params.get("title")
        group = self.request.query_params.get("group")
        price_min = self.request.query_params.get("price_min")
        price_max = self.request.query_params.get("price_max")

        if title:
            qs = qs.filter(title__icontains=title)
        if group:
            qs = qs.filter(group_id=group)
        if price_min:
            qs = qs.filter(price__gte=price_min)
        if price_max:
            qs = qs.filter(price__lte=price_max)

        return qs

    @action(detail=False, url_path="stats", methods=["GET"])
    def stats(self, request, *args, **kwargs):
        d = self.get_queryset().aggregate(
            count=Count("id"),
            avg=Avg("price"),
            min=Min("price"),
            max=Max("price"),
        )
        return Response({
            "count": d.get("count") or 0,
            "avg": d.get("avg") or 0,
            "min": d.get("min") or 0,
            "max": d.get("max") or 0,
        })

    @action(detail=False, url_path="export-excel", methods=["GET"])
    def export_excel(self, request, *args, **kwargs):
        if Workbook is None:
            return Response({"detail": "openpyxl не установлен в этом окружении"}, status=500)

        qs = self.get_queryset()
        wb = Workbook()
        ws = wb.active
        ws.title = "Меню"
        ws.append(["id", "title", "group", "price", "description"])

        for m in qs:
            ws.append([m.id, m.title, m.group_id, str(m.price), m.description or ""])

        resp = HttpResponse(
            content_type="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet"
        )
        resp["Content-Disposition"] = 'attachment; filename="menu.xlsx"'
        wb.save(resp)
        return resp

    @action(detail=False, url_path="export-word", methods=["GET"])
    def export_word(self, request, *args, **kwargs):
        if Document is None:
            return Response({"detail": "python-docx не установлен в этом окружении"}, status=500)

        qs = self.get_queryset()
        doc = Document()
        doc.add_heading("Меню", level=1)

        table = doc.add_table(rows=1, cols=5)
        hdr = table.rows[0].cells
        hdr[0].text = "id"
        hdr[1].text = "title"
        hdr[2].text = "group"
        hdr[3].text = "price"
        hdr[4].text = "description"

        for m in qs:
            row = table.add_row().cells
            row[0].text = str(m.id)
            row[1].text = m.title
            row[2].text = str(m.group_id or "")
            row[3].text = str(m.price)
            row[4].text = m.description or ""

        resp = HttpResponse(
            content_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
        )
        resp["Content-Disposition"] = 'attachment; filename="menu.docx"'
        doc.save(resp)
        return resp

class CustomerViewSet(ModelViewSet):
    queryset = Customer.objects.all().order_by("-id")
    serializer_class = CustomerSerializer
    permission_classes = [permissions.IsAuthenticated]

    def get_queryset(self):
        qs = super().get_queryset()
        name = self.request.query_params.get("name")
        phone = self.request.query_params.get("phone")

        if name:
            qs = qs.filter(name__icontains=name)
        if phone:
            qs = qs.filter(phone__icontains=phone)

        return qs

    @action(detail=False, url_path="stats", methods=["GET"])
    def stats(self, request, *args, **kwargs):
        d = self.get_queryset().aggregate(total=Count("id"))
        return Response({"total": d.get("total") or 0})

    @action(detail=False, url_path="export-excel", methods=["GET"])
    def export_excel(self, request, *args, **kwargs):
        if Workbook is None:
            return Response({"detail": "openpyxl не установлен в этом окружении"}, status=500)

        qs = self.get_queryset()
        wb = Workbook()
        ws = wb.active
        ws.title = "Клиенты"
        ws.append(["id", "name", "phone"])

        for c in qs:
            ws.append([c.id, c.name, c.phone or ""])

        resp = HttpResponse(
            content_type="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet"
        )
        resp["Content-Disposition"] = 'attachment; filename="customers.xlsx"'
        wb.save(resp)
        return resp

    @action(detail=False, url_path="export-word", methods=["GET"])
    def export_word(self, request, *args, **kwargs):
        if Document is None:
            return Response({"detail": "python-docx не установлен в этом окружении"}, status=500)

        qs = self.get_queryset()
        doc = Document()
        doc.add_heading("Клиенты", level=1)

        table = doc.add_table(rows=1, cols=3)
        hdr = table.rows[0].cells
        hdr[0].text = "id"
        hdr[1].text = "name"
        hdr[2].text = "phone"

        for c in qs:
            row = table.add_row().cells
            row[0].text = str(c.id)
            row[1].text = c.name
            row[2].text = c.phone or ""

        resp = HttpResponse(
            content_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
        )
        resp["Content-Disposition"] = 'attachment; filename="customers.docx"'
        doc.save(resp)
        return resp

class OrderViewSet(ModelViewSet):
    queryset = Order.objects.all().order_by("-id")
    serializer_class = OrderSerializer
    permission_classes = [permissions.IsAuthenticated]

    def get_queryset(self):
        qs = Order.objects.all().order_by("-id")

        customer = self.request.query_params.get("customer")
        status = self.request.query_params.get("status")

        if customer:
            qs = qs.filter(customer_id=customer)
        if status:
            qs = qs.filter(status=status)

        if not self.request.user.is_staff:
            qs = qs.filter(user=self.request.user)

        return qs

    @action(detail=False, url_path="stats", methods=["GET"])
    def stats(self, request, *args, **kwargs):
        qs = self.get_queryset()

        total_orders = qs.count()

        items_qs = OrderItem.objects.filter(order__in=qs)

        items_total = items_qs.count()
        qty_total = items_qs.aggregate(q=Sum("qty")).get("q") or 0

        revenue_expr = ExpressionWrapper(
            F("qty") * F("menu__price"),
            output_field=DecimalField(max_digits=12, decimal_places=2),
        )
        revenue = items_qs.aggregate(s=Sum(revenue_expr)).get("s") or 0

        return Response({
            "total_orders": total_orders,
            "items_total": items_total,
            "qty_total": qty_total,
            "revenue": revenue,
        })

    @action(detail=False, url_path="export-excel", methods=["GET"])
    def export_excel(self, request, *args, **kwargs):
        if Workbook is None:
            return Response({"detail": "openpyxl не установлен в этом окружении"}, status=500)

        qs = self.get_queryset()
        wb = Workbook()
        ws = wb.active
        ws.title = "Заказы"
        ws.append(["id", "customer", "status"])

        for o in qs:
            ws.append([o.id, o.customer_id, o.status])

        resp = HttpResponse(
            content_type="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet"
        )
        resp["Content-Disposition"] = 'attachment; filename="orders.xlsx"'
        wb.save(resp)
        return resp

    @action(detail=False, url_path="export-word", methods=["GET"])
    def export_word(self, request, *args, **kwargs):
        if Document is None:
            return Response({"detail": "python-docx не установлен в этом окружении"}, status=500)

        qs = self.get_queryset()
        doc = Document()
        doc.add_heading("Заказы", level=1)

        table = doc.add_table(rows=1, cols=3)
        hdr = table.rows[0].cells
        hdr[0].text = "id"
        hdr[1].text = "customer"
        hdr[2].text = "status"

        for o in qs:
            row = table.add_row().cells
            row[0].text = str(o.id)
            row[1].text = str(o.customer_id or "")
            row[2].text = o.status

        resp = HttpResponse(
            content_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
        )
        resp["Content-Disposition"] = 'attachment; filename="orders.docx"'
        doc.save(resp)
        return resp

class OrderItemViewSet(ModelViewSet):
    queryset = OrderItem.objects.all().order_by("-id")
    serializer_class = OrderItemSerializer
    permission_classes = [permissions.IsAuthenticated]

    def get_queryset(self):
        qs = OrderItem.objects.all().order_by("-id")

        order = self.request.query_params.get("order")
        menu = self.request.query_params.get("menu")
        qty_min = self.request.query_params.get("qty_min")
        qty_max = self.request.query_params.get("qty_max")

        if order:
            qs = qs.filter(order_id=order)
        if menu:
            qs = qs.filter(menu_id=menu)
        if qty_min:
            qs = qs.filter(qty__gte=qty_min)
        if qty_max:
            qs = qs.filter(qty__lte=qty_max)

        if not self.request.user.is_staff:
            qs = qs.filter(order__user=self.request.user)

        return qs

    @action(detail=False, url_path="stats", methods=["GET"])
    def stats(self, request, *args, **kwargs):
        d = self.get_queryset().aggregate(
            total=Count("id"),
            min_qty=Min("qty"),
            max_qty=Max("qty"),
            avg_qty=Avg("qty"),
        )
        return Response({
            "total": d.get("total") or 0,
            "min_qty": d.get("min_qty") or 0,
            "max_qty": d.get("max_qty") or 0,
            "avg_qty": d.get("avg_qty") or 0,
        })

    @action(detail=False, url_path="export-excel", methods=["GET"])
    def export_excel(self, request, *args, **kwargs):
        if Workbook is None:
            return Response({"detail": "openpyxl не установлен в этом окружении"}, status=500)

        qs = self.get_queryset()
        wb = Workbook()
        ws = wb.active
        ws.title = "Позиции заказов"
        ws.append(["id", "order", "menu", "qty"])

        for it in qs:
            ws.append([it.id, it.order_id, it.menu_id, it.qty])

        resp = HttpResponse(
            content_type="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet"
        )
        resp["Content-Disposition"] = 'attachment; filename="order_items.xlsx"'
        wb.save(resp)
        return resp

    @action(detail=False, url_path="export-word", methods=["GET"])
    def export_word(self, request, *args, **kwargs):
        if Document is None:
            return Response({"detail": "python-docx не установлен в этом окружении"}, status=500)

        qs = self.get_queryset()
        doc = Document()
        doc.add_heading("Позиции заказов", level=1)

        table = doc.add_table(rows=1, cols=4)
        hdr = table.rows[0].cells
        hdr[0].text = "id"
        hdr[1].text = "order"
        hdr[2].text = "menu"
        hdr[3].text = "qty"

        for it in qs:
            row = table.add_row().cells
            row[0].text = str(it.id)
            row[1].text = str(it.order_id or "")
            row[2].text = str(it.menu_id or "")
            row[3].text = str(it.qty)

        resp = HttpResponse(
            content_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
        )
        resp["Content-Disposition"] = 'attachment; filename="order_items.docx"'
        doc.save(resp)
        return resp
